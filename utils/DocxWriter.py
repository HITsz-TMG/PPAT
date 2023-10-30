from docx import Document
from docx.shared import RGBColor
from enum import Enum
from docx.enum.text import WD_COLOR_INDEX

'''r
文档指令: 
    添加标题: add_heading(text, level=1, style(optional))
    添加段落: add_content(text, style(optional))
    下一页: add_page_break()
    改变样式: set(start=None, end=None, pos=None, order_id=None, style(optional)):
    写入: write(order_id=None):
    保存关闭: close()
样式指令:
    color:文本颜色 (_font_color_translate)
    underline:是否有下划线
    bold:是否加粗 
    highlight:背景色 (_highlight_color_translate)
'''


class docx_writer:
    class docx_word:
        def __init__(self, word: str, style_dir=None):
            self.ch = word
            self.default_style = {
                'color': (0, 0, 0),
                'underline': False,
                'bold': False,
                'highlight': WD_COLOR_INDEX.AUTO
            }
            self.color = self.default_style['color']
            self.underline = self.default_style['underline']
            self.bold = self.default_style['bold']
            self.highlight = self.default_style['highlight']
            self.set(style_dir)

        def set(self, style_dir=None):
            # for reset
            if style_dir is None:
                style_dir = self.default_style

            self.color = self._font_color_translate(style_dir.get('color')) or self.color
            self.underline = style_dir.get('underline') or self.underline
            self.bold = style_dir.get('bold') or self.bold
            self.highlight = self._highlight_color_translate(style_dir.get('highlight')) or self.highlight

        def write(self, p):
            run = p.add_run(self.ch)
            run.font.color.rgb = RGBColor(self.color[0], self.color[1], self.color[2])
            run.underline = self.underline
            run.bold = self.bold
            run.font.highlight_color = self.highlight

        @staticmethod
        def _font_color_translate(color):
            if color is None: return None
            if isinstance(color, str):
                if color in ['red', '红']:
                    color = (255, 0, 0)
                elif color in ['pink', '粉红', '粉']:
                    color = (255, 105, 180)
                elif color in ['magenta', '洋红']:
                    color = (255, 0, 255)
                elif color in ['darkvoilet', 'purple', '紫']:
                    color = (148, 0, 211)
                elif color in ['indigo', '靛青']:
                    color = (75, 0, 130)
                elif color in ['mediumslateblue', '蓝灰']:
                    color = (123, 104, 238)
                elif color in ['blue', '蓝']:
                    color = (0, 0, 255)
                elif color in ['darkblue', '深蓝']:
                    color = (0, 0, 139)
                elif color in ['doderblue', '道奇蓝']:
                    color = (30, 144, 255)
                elif color in ['deepskyblue', '深天蓝', '天蓝']:
                    color = (0, 191, 255)
                elif color in ['cyan', '青']:
                    color = (0, 255, 255)
                elif color in ['mediumaquamarine', '碧绿']:
                    color = (0, 250, 154)
                elif color in ['springgreen', '春绿']:
                    color = (0, 191, 255)
                elif color in ['lime', '酸橙']:
                    color = (0, 255, 0)
                elif color in ['green', '绿']:
                    color = (0, 128, 0)
                elif color in ['yellow', '黄']:
                    color = (255, 255, 0)
                elif color in ['orangered', '橙', '橙红']:
                    color = (255, 69, 0)
                elif color in ['lightgrey', '浅灰', '灰']:
                    color = (211, 211, 211)
                elif color in ['darkgray', '深灰']:
                    color = (169, 169, 169)
                elif color in ['dimgray', '暗灰']:
                    color = (105, 105, 105)
                elif color in ['black', '黑']:
                    color = (0, 0, 0)
                elif color in ['white', '白']:
                    color = (255, 255, 255)
                else:
                    raise KeyError
            return color

        @staticmethod
        def _highlight_color_translate(color):
            if color is None: return None
            if color in ['default', '默认']:
                color = WD_COLOR_INDEX.AUTO
            elif color in ['black', '黑']:
                color = WD_COLOR_INDEX.BLACK
            elif color in ['blue', '蓝', ]:
                color = WD_COLOR_INDEX.BLUE
            elif color in ['lightgreen', '浅绿']:
                color = WD_COLOR_INDEX.BRIGHT_GREEN
            elif color in ['gold', '金']:
                color = WD_COLOR_INDEX.DARK_YELLOW
            elif color in ['lightgray', '浅灰']:
                color = WD_COLOR_INDEX.GRAY_25
            elif color in ['gray', '灰']:
                color = WD_COLOR_INDEX.GRAY_50
            elif color in ['green', '绿']:
                color = WD_COLOR_INDEX.GREEN
            elif color in ['pink', '粉']:
                color = WD_COLOR_INDEX.PINK
            elif color in ['red', '红']:
                color = WD_COLOR_INDEX.RED
            elif color in ['teal', '青']:
                color = WD_COLOR_INDEX.TEAL
            elif color in ['qurquouse', '青绿']:
                color = WD_COLOR_INDEX.TURQUOISE
            elif color in ['violet', '紫']:
                color = WD_COLOR_INDEX.VIOLET
            elif color in ['white', '白']:
                color = WD_COLOR_INDEX.WHITE
            elif color in ['yellow', '黄']:
                color = WD_COLOR_INDEX.YELLOW
            else:
                raise KeyError
            return color

    def __init__(self, path='./temp.docx'):
        self.doc = Document()
        self.path = path
        self.order_list = []
        self.order_start = 0

    def _text_trans(self, text, style_dir):
        if isinstance(text, str):
            return [self.docx_word(text, style_dir)]
        else:
            return [self.docx_word(i, style_dir) for i in text]

    def add_heading(self, text, level=1, **kwargs):
        text = self._text_trans(text, kwargs)
        self.order_list.append(['heading', level, text])
        return len(self.order_list) - 1

    def add_content(self, text, **kwargs):
        text = self._text_trans(text, kwargs)
        self.order_list.append(['content', None, text])
        return len(self.order_list) - 1

    def add_page_break(self):
        self.order_list.append(['page break', None, None])
        return len(self.order_list) - 1

    def set(self, start=None, end=None, pos=None, order_id=None, **kwargs):
        if order_id is None: order_id = len(self.order_list) - 1
        text = self.order_list[order_id][2]
        if isinstance(pos, int):
            text[pos].set(kwargs)
        else:
            if pos is None:
                if start is None: start = 0
                if end is None: end = len(self.order_list[order_id])
                pos = list(range(start, end))
            for p in pos:
                text[p].set(kwargs)

    def _run(self, order):
        cmd, arg, text = order
        if cmd == 'heading':
            p = self.doc.add_heading('', level=arg)
            for word in text:
                word.write(p)
        elif cmd == 'page break':
            self.doc.add_page_break()
        elif cmd == 'content':
            p = self.doc.add_paragraph('')
            for word in text:
                word.write(p)

    def write(self, order_id=None):
        if order_id is None: order_id = len(self.order_list) - 1
        for i in range(self.order_start, order_id + 1):
            self._run(self.order_list[i])
        self.order_start = order_id + 1

    def close(self):
        self.write()
        self.doc.save(self.path)


if __name__ == '__main__':
    doc = Document()
